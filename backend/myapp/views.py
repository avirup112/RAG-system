import json
from django.http import JsonResponse, FileResponse
from rest_framework.decorators import api_view, parser_classes
from rest_framework.parsers import MultiPartParser
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
import os, pickle, faiss, numpy as np, PyPDF2, docx
from groq import Groq
from transformers import pipeline
from .models import *
from sentence_transformers import SentenceTransformer, util
from .prompts import LEGAL_ASSISTANT_PROMPT
import re
from langdetect import detect
from docx import Document as DocxDocument
import io 
from fpdf import FPDF
import pytesseract
from PIL import Image
import pytesseract
import spacy



embed_model = SentenceTransformer("distiluse-base-multilingual-cased-v2")
spacy_en = spacy.load("en_core_web_sm")
INDEX_DIR = "rag_indexes"
os.makedirs(INDEX_DIR, exist_ok=True)
ner_pipeline = pipeline("ner", model="dslim/bert-base-NER", aggregation_strategy="simple")

groq_api_key = os.getenv("GROQ_API_KEY")
client = Groq(api_key=groq_api_key)

# Clause classifier
def classify_clause(header, body):
    known_types = {
        "termination": ["termination", "end of agreement", "end date"],
        "confidentiality": ["confidentiality", "non-disclosure", "nda"],
        "governing_law": ["jurisdiction", "governing law", "legal provisions"],
        "liability": ["liability", "indemnification", "damages"],
        "dispute_resolution": ["dispute", "arbitration", "mediation"],
        "payment": ["payment", "fees", "charges"],
        "intellectual_property": ["intellectual property", "ownership", "IP rights"]
    }

    clause_embedding = embed_model.encode(f"{header} {body}", convert_to_tensor=True)

    scores = {}
    for label, keywords in known_types.items():
        keyword_embeds = embed_model.encode(keywords, convert_to_tensor=True)
        score = util.cos_sim(clause_embedding, keyword_embeds).max().item()
        scores[label] = score

    best_match = max(scores, key=scores.get)
    return best_match if scores[best_match] > 0.4 else "unlabeled"

def extract_text(file, filename):
    if filename.endswith(".pdf"):
        reader = PyPDF2.PdfReader(file)
        return "\n".join([p.extract_text() or "" for p in reader.pages])
    elif filename.endswith(".docx"):
        docx_file = docx.Document(file)
        return "\n".join([p.text for p in docx_file.paragraphs])
    elif filename.endswith(".txt"):
        return file.read().decode("utf-8")
    return ""

def chunk_text(text, max_len=1000):
    return [text[i:i+max_len] for i in range(0, len(text), max_len)]

# Clause Extraction View
@csrf_exempt
def extract_clauses(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST request required'}, status=405)

    try:
        content = json.loads(request.body)
        text = content.get("text")
        if not text:
            return JsonResponse({'error': 'Missing text in request'}, status=400)

        lang = detect(text)

        clauses = []
        current_clause = {"header": None, "body": ""}

        lines = text.splitlines()
        for line in lines:
            if re.match(r"^([0-9.]+)?\s*[A-Z][A-Za-z\s,/\\-]{2,}$", line.strip()):
                if current_clause["header"]:
                    clauses.append(current_clause)
                current_clause = {"header": line.strip(), "body": ""}
            else:
                current_clause["body"] += line.strip() + " "

        if current_clause["header"]:
            clauses.append(current_clause)

        results = []
        for clause in clauses:
            header = clause["header"]
            body = clause["body"]
            clause_type = classify_clause(header, body)

            # Save to DB
            Clause.objects.create(
                header=header,
                body=body,
                clause_type=clause_type
            )

            entities = [
                {"word": ent["word"], "label": ent["entity_group"], "score": round(ent["score"], 3)}
                for ent in ner_pipeline(body)
            ]

            results.append({
                "header": header,
                "body": body.strip(),
                "clause_type": clause_type,
                "language": lang,
                "entities": entities
            })

        return JsonResponse({"clauses": results}, safe=False)

    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@api_view(['POST'])
def ask_legal_bot(request):
    user_query = request.data.get("query")
    doc_title = request.data.get("doc_title")
    screenshot = request.FILES.get("screenshot")
    if not user_query or not doc_title or not screenshot:
        return JsonResponse({"error": "No query provided"}, status=400)
    
    combined_chunks = []
    
    if doc_title:
        for title in doc_title:
            try:
                doc_obj = UploadedDocument.objects.get(title=doc_title)
                index = faiss.read_index(doc_obj.index_path)
                with open(doc_obj.chunks_path, "rb") as f:
                    chunks = pickle.load(f)
                
                query_emb = embed_model.encode([user_query])
                D, I = index.search(np.array(query_emb), k=3)
                top_chunks = [chunks[i] for i in I[0]]
                combined_chunks.extend(top_chunks)
            except UploadedDocument.DoesNotExist:
                return ask_legal_bot(request)
    if screenshot:
        extracted_text = extract_text_from_image(screenshot)
        if len(extracted_text.strip()) > 50:  # Only if meaningful text
            combined_chunks.append(extracted_text)
    
    if not combined_chunks:
        prompt = LEGAL_ASSISTANT_PROMPT.format(user_query=user_query)
        response = client.chat.completions.create(
            model="llama3-70b-8192",
            messages=[
                {"role": "system", "content": "You are a legal assistant trained in Indian law, constitution, IPC, and CrPC."},
                {"role": "user", "content": prompt}
            ]
        )
        return JsonResponse({"response": response.choices[0].message.content})
    
    context = "\n\n".join(combined_chunks)
    prompt = (
        "You are a legal assistant. Based on the user's query, answer using the following document content:\n\n"
        f"Query: {user_query}\n\n"
        f"Relevant Content:\n{context}"
    )
    
    response = client.chat.completions.create(
        model="llama3-70b-8192",
        messages=[
            {"role": "system", "content": "You answer legal questions based on uploaded contract contents."},
            {"role": "user", "content": prompt}
        ]
    )

    return JsonResponse({"response": response.choices[0].message.content})
        
@api_view(['POST'])
@parser_classes([MultiPartParser])
def summarize_document(request):
    file = request.FILES.get("file")
    user_query = request.data.get("user_query") or "Please simplify the legal document."

    if not file:
        return JsonResponse({"error": "No file uploaded."}, status=400)

    filename = file.name.lower()
    name_base = os.path.splitext(filename)[0]
    index_path = os.path.join(INDEX_DIR, f"{name_base}.index")
    chunks_path = os.path.join(INDEX_DIR, f"{name_base}_chunks.pkl")

    try:
        doc_obj = UploadedDocument.objects.get(title=name_base)
        index = faiss.read_index(doc_obj.index_path)
        with open(doc_obj.chunks_path, "rb") as f:
            chunks = pickle.load(f)
    except UploadedDocument.DoesNotExist:
        # Extract text and create new embedding
        full_text = extract_text(file, filename)
        if not full_text.strip():
            return JsonResponse({"error": "The file is empty or couldn't be read."}, status=400)

        chunks = chunk_text(full_text)
        embeddings = embed_model.encode(chunks)

        dim = embeddings[0].shape[0]
        index = faiss.IndexFlatL2(dim)
        index.add(np.array(embeddings))

        faiss.write_index(index, index_path)
        with open(chunks_path, "wb") as f:
            pickle.dump(chunks, f)

        UploadedDocument.objects.create(
            title=name_base,
            full_text=full_text,
            index_path=index_path,
            chunks_path=chunks_path
        )

    # Query embedding + retrieve top chunks
    query_emb = embed_model.encode([user_query])
    D, I = index.search(np.array(query_emb), k=5)
    top_chunks = [chunks[i] for i in I[0]]
    combined_context = "\n\n".join(top_chunks)

    prompt = (
        "You are a legal assistant. Based on the user's query, simplify or summarize the following content:\n\n"
        f"User Query: {user_query}\n\n"
        f"Relevant Content:\n{combined_context}"
    )

    response = client.chat.completions.create(
        model="llama3-70b-8192",
        messages=[
            {"role": "system", "content": "You simplify legal documents for laypersons."},
            {"role": "user", "content": prompt}
        ]
    )

    return JsonResponse({"summary": response.choices[0].message.content})


@csrf_exempt
@require_http_methods(["GET"])
def filter_clauses(request):
    clause_type = request.GET.get("type")
    language = request.GET.get("lang")

    clauses = Clause.objects.all()
    if clause_type:
        clauses = clauses.filter(clause_type=clause_type)
    if language:
        clauses = [c for c in clauses if detect(c.body) == language]

    result = [
        {
            "header": c.header,
            "body": c.body,
            "clause_type": c.clause_type,
        }
        for c in clauses
    ]

    return JsonResponse({"filtered_clauses": result})

@api_view(['POST'])
def generate_legal_template(request):
    data = request.data

    # Collect user input fields
    document_type = data.get("document_type")
    party_one = data.get("party_one")
    party_two = data.get("party_two")
    effective_date = data.get("effective_date")
    duration = data.get("duration")
    payment_terms = data.get("payment_terms")
    jurisdiction = data.get("jurisdiction")
    user_name = data.get("user_name", "Anonymous")
    extra_details = data.get("extra_details", "")
    download_format = data.get("download_format", "txt")

    if not all([document_type, party_one, party_two, effective_date]):
        return JsonResponse({"error": "Missing required fields."}, status=400)

    # Construct prompt with details
    prompt = f"""
    You are a legal expert. Draft a {document_type} based on the following details:
    - Party One: {party_one}
    - Party Two: {party_two}
    - Effective Date: {effective_date}
    - Duration: {duration or 'Not specified'}
    - Payment Terms: {payment_terms or 'Not specified'}
    - Jurisdiction: {jurisdiction or 'Not specified'}
    - Additional Details: {extra_details or 'None'}
    Use proper legal language and structure.
    """

    # Generate response
    response = client.chat.completions.create(
        model="llama3-70b-8192",
        messages=[
            {"role": "system", "content": "You are a legal document generator trained in contract and clause drafting."},
            {"role": "user", "content": prompt}
        ]
    )

    generated_text = response.choices[0].message.content.strip()

    # Save in DB
    doc = LegalDocument.objects.create(
        title=f"{document_type} - {party_one} vs {party_two}",
        content=generated_text,
        document_type=document_type,
        user_name=user_name,
        download_format=download_format
    )

    # Generate download
    if download_format == "docx":
        buffer = save_as_docx(generated_text)
        return FileResponse(buffer, as_attachment=True, filename=f"{doc.title}.docx")
    elif download_format == "pdf":
        buffer = save_as_pdf(generated_text)
        return FileResponse(buffer, as_attachment=True, filename=f"{doc.title}.pdf")

    return JsonResponse({"template": generated_text})

def save_as_docx(text):
    doc = DocxDocument()
    doc.add_paragraph(text)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def save_as_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for line in text.split("\n"):
        pdf.multi_cell(0, 10, line)
    buffer = io.BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer

@api_view(['POST'])
def semantic_clause_search(request):
    query = request.data.get("query")
    if not query:
        return JsonResponse({"error": "Query required."}, status=400)

    query_embedding = embed_model.encode(query, convert_to_tensor=True)
    all_clauses = Clause.objects.all()

    results = []
    for clause in all_clauses:
        text = f"{clause.header} {clause.body}"
        clause_embedding = embed_model.encode(text, convert_to_tensor=True)
        score = util.cos_sim(query_embedding, clause_embedding).item()

        if score > 0.4:
            results.append({
                "header": clause.header,
                "body": clause.body,
                "clause_type": clause.clause_type,
                "similarity": score
            })

    results = sorted(results, key=lambda x: x["similarity"], reverse=True)

    return JsonResponse({"results": results})


def detect_and_redact_pii(text, redact=True):
    doc = spacy_en(text)
    redacted_text = text
    
    pii_entities = []
    for ent in doc.ents:
        if ent.label_ in ["PERSON", "ORG", "DATE", "TIME", "PER", "GPE"]:
            pii_entities.append((ent.text, ent.label_))
            if redact:
                redacted_text = redacted_text.replace(ent.text, "[REDACTED]")
                
    patterns ={
        "PHONE": r"\b[789]\d{9}\b",
        "AADHAAR": r"\b\d{4}\s\d{4}\s\d{4}\b",
        "PAN": r"\b[A-Z]{5}[0-9]{4}[A-Z]\b",
        "EMAIL": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
    }
    for label, pattern in patterns.items():
        matches = re.findall(pattern, text)
        for match in matches:
            pii_entities.append((match, label))
            if redact:
                redacted_text = redacted_text.replace(match, "[REDACTED]")
    
    return {
        "pii_entities": pii_entities,
        "redacted_text": redacted_text
    }
    
@api_view(["POST"])
def detect_pii(request):
    text = request.data.get("text")
    redact = request.data.get("redact", True)
    
    if not text:
        return JsonResponse({"error": "Missing Text is request"}, status=400)
    
    result = detect_and_redact_pii(text, redact=redact)
    
    return JsonResponse(result)

@api_view(['POST'])
@parser_classes([MultiPartParser])
def chat_from_screenshot(request):
    image = request.FILES.get('image')
    user_query = request.data.get("query")
    
    if not image or not user_query:
        return JsonResponse({"error": "Missing screenshot or query"}, status=400)
    
    img = Image.open(image)
    extracted_text = pytesseract.image_to_string(img)
    
    if not extracted_text.strip():
        return JsonResponse({"error": "Unable to extract text from image"}, status=400)
    
    prompt = (
        f"You are a legal assistant. Based on the extracted legal document content below, answer the question:\n\n"
        f"Extracted Content:\n{extracted_text}\n\n"
        f"User Query: {user_query}"
    )
    
    response = client.chat.completions.create(
        model="llama3-70b-8192",
        messages=[
            {"role": "system", "content": "You answer legal queries based on document content."},
            {"role": "user", "content": prompt}
        ]
    )
    
    return JsonResponse({"response": response.choices[0].message.content})
    
    